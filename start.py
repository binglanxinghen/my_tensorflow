import datetime
from flask import Flask
from flask import render_template, redirect, url_for, request, session
from werkzeug import secure_filename
from flask import send_file, send_from_directory
from flask import make_response
import os
from flask_paginate import Pagination,get_page_parameter
import time
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash  # 密码保护，使用hash方法
from flask_sqlalchemy import SQLAlchemy

from docx import Document
from docx.shared import Inches
from PIL import Image
import json
from predict import YoloTest
yolo_mini = YoloTest()
# import python4
basedir = os.path.abspath(os.path.dirname(__file__)) #定义一个根目录 用于保存图片用

app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///students.sqlite3'
app.config['SECRET_KEY'] = "random string"

db = SQLAlchemy(app)

#user表
class User(db.Model):
    __tablename__ = 'user'
    id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    username = db.Column(db.String(20), nullable=False)
    _password = db.Column(db.String(200), nullable=False)  # 内部使用

    @property
    def password(self):  # 定义一个外部使用的密码
        return self._password

    @password.setter  # 设置密码加密
    def password(self, row_password):
        self._password = generate_password_hash(row_password)

    # 检查密码是否正确
    def check_password(self, row_password):  # 定义一个反向解密的函数
        result = check_password_hash(self._password, row_password)
        return result

#
class History(db.Model):
    __tablename__ = 'history'
    history_id = db.Column(db.Integer, primary_key=True, autoincrement=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    result = db.Column(db.Text, nullable=False)
    file_name = db.Column(db.String(225),nullable=False)
    time = db.Column(db.DateTime, default=datetime.now())
    # image_url = db.Column(db.String(225), nullable=False)
    # author = db.relationship('User', backref=db.backref('questions'))

# db.drop_all()
# print ("Table delete successfully")
db.create_all()
print ("Table created successfully")

def count_classes(flow_list):
	classes = {"daisy":0,"dandelion":0,"rose":0,"sunflowers":0,"tulips":0}
	sum=0;
	for item in flow_list:
		if item.split()[0] in classes.keys():
			classes[item.split()[0]]= classes[item.split()[0]] + 1
			sum=sum+1; 
	if(sum==0):
		s="Unkonown"
		return s;
	s=""
	for key,value in classes.items():
		if(value!=0):
			if(key=='daisy'):
				key='雏菊'
			if(key=='dandelion'):
				key='蒲公英'
			if(key=='rose'):
				key='玫瑰花'
			if(key=='sunflowers'):
				key='向日葵'
			if(key=='tulips'):
				key='郁金香'
			s=s+(key+":"+str(value)+"\n")
	return s;			
	# return "雏菊: "+str(classes['daisy'])+'\n'+"蒲公英: "+str(classes['dandelion'])+'\n'+"玫瑰花: "+str(classes['rose'])+'\n'+"向日葵: "+str(classes['sunflowers'])+'\n'+"郁金香: "+str(classes['tulips'])+'\n'

@app.route("/download/<history_id>", methods=["GET"])
def download(history_id):
    history = History.query.filter(History.history_id == history_id).first()
    images = "static/result/"+str(history.user_id)+"/"+history.file_name  # 保存在本地的图片
    doc = Document()
    doc.add_heading('结果报告', level=2)
    try:
        doc.add_picture(images, width=Inches(2))  # 添加图, 设置宽度
    except Exception:
        jpg_ima = Image.open(images)  # 打开图片
        jpg_ima.save('0.jpg')  # 保存新的图片
        doc.add_picture(images, width=Inches(2))  # 添加图, 设置宽度
    doc.add_paragraph("结果：")  # 添加文字
    doc.add_paragraph(history.result)  # 添加文字
    doc.add_paragraph("时间：" + str(history.time))  # 添加文字
    doc.add_paragraph("本报告来自XXX网站")  # 添加文字
    doc.save('static/result.docx')  # 保存路径
    directory = os.getcwd()  # 假设在当前目录
    response = make_response(send_from_directory(directory+"/static","result.docx" , as_attachment=True))
    response.headers["Content-Disposition"] = "attachment; filename={}".format("result.docx" .encode().decode('latin-1'))
    return response

#调用predict
@app.route("/forecast", methods=["POST"])
def forecast(path,filename):
    result_str=""
    # results = os.system( "python predict.py %s %s %s %s" % ("-i","./static/img/"+imgName,"-o","./static/result/"+imgName))
    if not session.get('id'):
        #游客
        yolo_mini.evaluate("./" + path + filename, "./" + path + filename)
        # results = os.system(
        #     "python predict.py %s %s %s %s" % ("-i", "./" + path + filename, "-o", "./" + path + filename))
    else:
        #登录用户
        yolo_mini.evaluate( "./" + path + filename, "./static/result/" + str(session.get('id')) + "/" + filename)
        # results = os.system(
        #     "python predict.py %s %s %s %s" % ("-i", "./" + path + filename, "-o", "./static/result/" + str(session.get('id')) + "/" + filename))
    with open("flow_result.txt",'r') as f:
        lines = f.readlines()
        result_str = count_classes(lines)
    return result_str

#主页
@app.route('/')
def index():
    return render_template('index.html')

# @app.route('/')
# def index():
#     page=request.args.get('page',1,type=int)
#     pagination=History.query.order_by(History.time.desc()).paginate(page,per_page=3,error_out=False)
#     users=pagination.items
#     return render_template('test.html', historys=users,pagination=pagination)

@app.route('/history')
def history():
    if not session.get('id'):
        return render_template('login.html')
    id=session.get('id')
    page = request.args.get('page', 1, type=int)
    pagination = History.query.filter(History.user_id == id).order_by(History.time.desc()).paginate(page, per_page=3, error_out=False)
    historys = pagination.items
    # context = {
    #     'historys':historys,
    #     "pagination":pagination
    # }
    return render_template('history.html', historys=historys, pagination=pagination,base_url="result/"+str(id)+"/")
    # return render_template('history.html',**context)

def add_history(user_id,fname,result):
    history = History(result=result, file_name=fname, user_id=user_id)
    db.session.add(history)
    db.session.flush()
    hid = history.history_id
    db.session.commit()
    db.session.flush()
    return hid

#删除历史纪录
@app.route('/delete/<history_id>')
def del_history(history_id):
    if not session.get('id'):
        return redirect(url_for('login'))
    else:
        history = History.query.filter(History.history_id == history_id).first()
        if session.get('id')!=history.user_id:
            return redirect(url_for('login'))
        else:
            db.session.delete(history)
            db.session.commit()
            db.session.flush()
        return redirect(url_for('history'))

#拍照上传
@app.route('/camera',methods=['GET','POST'])
def camera():
    print('camera')
    if request.method == 'POST':
        print('camera')
        f = request.files['file1']
        basepath = os.path.dirname(__file__)  # 当前文件所在路径
        fname= time.strftime("%Y%m%d%H%M%S",time.localtime(time.time()))+secure_filename(f.filename)
        if not session.get('id'):
            # 游客
            print('游客')
            upload_path = os.path.join(basepath, 'static/image/', fname)
            # 游客的放在static/image/
            f.save(upload_path)
            print(upload_path)
            res = forecast('static/image/', fname)
            return render_template('detail.html',image_url="image/"+fname,result=res,time=datetime.now(),history=None)
        dirs = 'img/' + str(session.get('id'))+'/'
        path = 'static/'+dirs
        # print(path)
        if not os.path.exists(path):
            print(path)
            os.makedirs(path)
        upload_path = os.path.join(basepath, path, fname)
        # 保存文件位置static/img/<user_id>/
        f.save(upload_path)
        #获取结果
        result = forecast(path,fname)
        hid = add_history(session.get('id'), fname, result)
        # print(hid)
        return redirect(url_for('detail', history_id=hid))
    return render_template('index.html')

#上传文件
@app.route('/upload',methods=['GET','POST'])
def upload():
    print('upload')
    if request.method == 'POST':
        print('upload')
        f=request.files['file2']
        basepath = os.path.dirname(__file__)  # 当前文件所在路径
        fname = time.strftime("%Y%m%d%H%M%S", time.localtime(time.time())) + secure_filename(f.filename)
        if not session.get('id'):
            # 游客
            print('游客')
            upload_path = os.path.join(basepath, 'static/image/', fname)
            # 游客的放在static/image/
            f.save(upload_path)
            print(upload_path)
            res = forecast('static/image/', fname)
            return render_template('detail.html', image_url="image/"+fname, result=res, time=datetime.now(), history=None)
        dirs = 'img/' + str(session.get('id')) + '/'
        path = 'static/' + dirs
        # print(path)
        if not os.path.exists(path):
            print(path)
            os.makedirs(path)
        upload_path = os.path.join(basepath, path, fname)
        # 保存文件位置static/img/<user_id>/
        f.save(upload_path)
        # 获取结果
        result = forecast(path, fname)
        hid = add_history(session.get('id'), fname, result)
        # print(hid)
        return redirect(url_for('detail', history_id=hid))
    return render_template('index.html')

#时间格式化
@app.template_filter('datetime_format')
def _jinja2_filter_datetime_format(datetimeValue, format='%y/%m/%d %H:%M:%S'):
    """convert a datetime to a different format."""
    print(datetime)
    return datetimeValue.strftime(format)

#详情
@app.route('/detail/<history_id>')
def detail(history_id):
    if not session.get('id'):
        return render_template('login.html')
    history = History.query.filter(History.history_id == history_id).first()
    return render_template('detail.html',image_url="result/"+str(history.user_id)+"/"+history.file_name, history = history)

#关于
@app.route('/about')
def about():
    return render_template('about.html')

#登录
@app.route('/login/',methods=['GET','POST'])
def login():
    if request.method=='GET':
        return render_template('login.html')
    else:
        username = request.form.get('username')
        password = request.form.get('password')
        # print(username)
        # print(password)
        user = User.query.filter(User.username == username).first()
        # print(user)
        if user:
            if user.check_password(password):
                session['user']=username
                session['id']=user.id
                session.permanent = True
                return redirect(url_for('index'))
            else:
                return render_template('login.html', error='user password error')
        else:
            return render_template('login.html',error="username not exited")

#退出
@app.route('/logout/',methods=['GET','POST'])
def logout():
    session.clear()
    return redirect(url_for('index'))

#注册
@app.route('/register',methods=['GET','POST'])
def register():
    if request.method=='GET':
        return render_template('register.html')
    else:
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter(User.username==username).first()
        if user:
            return render_template('register.html',error='username existed')
        else:
            user = User(username=username,password=password)
            db.session.add(user)
            db.session.flush()
            user_id = user.id
            db.session.commit()
            #创建文件夹
            path1 = 'static/result/'+str(user_id)
            path2 = 'static/img/' + str(user_id)
            os.makedirs(path1)
            os.makedirs(path2)
            return redirect(url_for('login'))

# 修改密码
@app.route('/edit_password/', methods=['GET', 'POST'])
def edit_password():
    if request.method == 'GET':
        return render_template("edit_password.html")
    else:
        newpassword = request.form.get('password')
        user = User.query.filter(User.username == request.form.get('username')).first()
        if user:
            user.password = newpassword
            db.session.commit()
        else:
            return 'not existed'
        return redirect(url_for('index'))
#花卉介绍
@app.route('/rose')
def rose():
     return render_template('rose.html')
@app.route('/daisy')
def daisy():
     return render_template('daisy.html')
@app.route('/sunflowers')
def sunflowers():
     return render_template('sunflowers.html')
@app.route('/tulips')
def tulips():
     return render_template('tulips.html')
@app.route('/dandelion')
def dandelion():
     return render_template('dandelion.html')






if __name__ =="__main__":
     app.run(host='0.0.0.0',port=80)
